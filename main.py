"""Script for identifying documents containing specific strings
"""

import os
import sys
import logging
import datetime
import shutil
import pypdf
import docx
import openpyxl
import ttk
from tkinter import messagebox
import tkinter as tk
import webbrowser
from dataclasses import dataclass
from openpyxl.styles import PatternFill, Alignment
from openpyxl.utils import get_column_letter

__version__="v0.1"

def log_config():
    """Performs a logging basic setup"""
    handler_to_file = logging.FileHandler("log.log", "w")
    handler_to_file.setLevel(logging.DEBUG)
    handler_to_console = logging.StreamHandler()
    handler_to_console.setLevel(logging.ERROR)
    logging.basicConfig(
        handlers=[
            handler_to_file,
            handler_to_console,
        ],
        format="%(asctime)s: %(levelname)s %(filename)s %(lineno)s: %(message)s",
        level=logging.DEBUG,
    )


def get_base_dir():
    """Get base dir, regardless if running from script or frozen exe"""
    # https://stackoverflow.com/questions/404744/determining-application-path-in-a-python-exe-generated-by-pyinstaller
    if getattr(sys, "frozen", False):
        logging.debug("Running from executable file")
        return os.path.dirname(sys.executable)
    else:
        logging.debug("Running from script")
        return os.path.dirname(os.path.realpath(__file__))

@dataclass
class Gui:
    """Run Button and Progress Bar"""
    base_dir : str

    def __post_init__(self):
        self.tk = tk.Tk()
        self.tk.title("Documents Filter "+__version__)
        self.tk.geometry("500x250")
        self.tk.iconbitmap(default=os.path.join(self.base_dir, "logo.ico"))
        self.full_bar = 480
        self.source_label = tk.Label(
            self.tk,
            text="See the code at github.com/lmponcio",
            cursor="hand1",
            font=("Courier", 8),
        )
        self.source_label.bind(
            "<Button-1>",
            lambda x: webbrowser.open_new(
                "https://github.com/lmponcio/DocumentsFilter"
            ),
        )
        self.source_label.pack(side=tk.BOTTOM, pady=5)
        self.progress = ttk.Progressbar(
            self.tk, orient=tk.HORIZONTAL, length=self.full_bar, mode="determinate"
        )
        self.progress.pack(side=tk.BOTTOM, pady=30)

    def advance(self, percentage):
        self.progress["value"] = percentage
        self.tk.update_idletasks()

    def success(self):
        self.success_label = tk.Label(
            self.tk, text="Finished!", bg="green", fg="white", width="32", height="2"
        )
        self.success_label.pack(side=tk.BOTTOM)

    def remove_success(self):
        self.success_label.destroy()
        self.tk.update_idletasks()


def main():
    if not getattr(sys, "frozen", False):
        log_config()
    logging.debug("Program started")
    base_dir = get_base_dir()
    gui = Gui(base_dir)
    tk.Button(
        gui.tk,
        text="Run Filters",
        command=lambda: run_filter(gui),
        width="30",
        height="2",
    ).pack(side=tk.TOP, pady=30)
    tk.mainloop()


def run_filter(gui):
    gui.advance(10)
    if hasattr(gui, "success_label"):
        gui.remove_success()
    base_dir = gui.base_dir
    logging.debug("Main folder found at %s", base_dir)
    doc_mgr = DocMgr(base_dir, gui)
    gui.advance(20)
    doc_mgr.filter_docs()


@dataclass
class Document:
    "Class representing file of any extension to scan"
    path: str

    def __post_init__(self):
        # All text of the document to be contained in a single string
        self.texts = ""
        self.filters_dict = {}

    def _set_filters_dict(self, filters):
        for filter in filters:
            self.filters_dict[filter] = False

    def passes_filters(self, filters):
        self._set_filters_dict(filters)
        for filter in filters:
            if filter in self.texts:
                self.filters_dict[filter] = True
        for partial_result in self.filters_dict.values():
            if partial_result is not True:
                return False
        return True

    def get_path(self):
        return self.path

    def get_texts(self):
        return self.texts


class PdfDoc(Document):
    "Class representing Pdf file to scan"

    def __init__(self, path):
        super().__init__(path)
        self.reader = pypdf.PdfReader(self.path)
        for page in self.reader.pages:
            this_text = page.extract_text()
            self.texts += " " + this_text.lower()


class DocxDoc(Document):
    "Class representing Docx file to scan"

    def __init__(self, path):
        super().__init__(path)
        self.reader = docx.Document(self.path)
        for para in self.reader.paragraphs:
            self.texts += " " + para.text.lower()
        for table in self.reader.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.texts += " " + para.text.lower()


@dataclass
class DocMgr:
    base_dir: str
    gui: Gui

    def __post_init__(self):
        self.docs = []
        self.filters = []
        self.passed_list = []
        self.input_dir_basename = "Files to Filter"
        self.input_dir = os.path.join(self.base_dir, self.input_dir_basename)
        self.filter_filename = "Filters.txt"
        self.filter_path = os.path.join(self.base_dir, self.filter_filename)
        self._check_setup()

    def _check_setup(self):
        if not os.path.exists(self.input_dir):
            logging.error("No input dir found at %s", self.input_dir)
            messagebox.showerror(
                title="Error",
                message=f'The program can\'t find the "{self.input_dir_basename}" folder.\n\nLocate "{self.input_dir_basename}" folder in the same folder as the program file and try again.\n\nDouble-check the folder name\'s spelling.',
            )
            sys.exit()
        if not os.path.exists(self.filter_path):
            logging.error("No filters file found at %s", self.filter_path)
            messagebox.showerror(
                title="Error",
                message=f'The program can\'t find "{self.filter_filename}".\n\nLocate "{self.filter_filename}" in the same folder as the program file and try again.\n\nDouble-check the file name\'s spelling.',
            )
            sys.exit()

    def _import_docs(self):
        """Imports documents from input dir"""
        for root, dirs, files in os.walk(self.input_dir):
            for file in files:
                file_path = os.path.join(root, file)
                if file.lower().endswith(".pdf"):
                    self.docs.append(PdfDoc(file_path))
                elif file.lower().endswith(".docx"):
                    self.docs.append(DocxDoc(file_path))
                else:
                    logging.error("Document with unexpected extension: %s", file_path)
        logging.debug("A total of %s docs were imported: %s", len(self.docs), self.docs)

    def _import_filters(self):
        """Imports filters from filters file"""
        with open(self.filter_path) as file:
            for line in file:
                text = line.strip().lower()
                self.filters.append(text)
        logging.debug("The following filters were imported: %s", self.filters)

    def _create_summary_sheet(self, dir):
        # Workbook and sheets conf
        wb = openpyxl.Workbook()
        ws = wb.active

        # Detail sheet
        ## Build Headers
        header_file = ws.cell(row=1, column=1, value="File")
        header_file.fill = PatternFill(start_color="729fcf", fill_type="solid")
        ws.column_dimensions["A"].width = 35
        col = 2
        for filter in self.filters:
            header_filter = ws.cell(row=1, column=col, value=filter)
            header_filter.fill = PatternFill(start_color="b4c7dc", fill_type="solid")
            header_filter.alignment = Alignment(horizontal="center")
            ws.column_dimensions[get_column_letter(col)].width = 20
            col += 1
        ## Fill Records
        for row in ws.iter_rows(min_row=1, max_row=1, min_col=1):
            headers = row
        row = 2
        for doc in self.docs:
            for cell in headers:
                col = cell.column
                if col == 1:
                    file_name = os.path.basename(doc.path)
                    ws.cell(row=row, column=col, value=file_name)
                else:
                    filter = ws.cell(row=1, column=col).value
                    if doc.filters_dict[filter]:
                        this_cell = ws.cell(row=row, column=col, value="YES")
                        this_cell.alignment = Alignment(horizontal="center")
                        this_cell.fill = PatternFill(
                            start_color="dde8cb", fill_type="solid"
                        )
            row += 1
        wb.save(os.path.join(dir, "Summary.xlsx"))

    def _copy_passed_docs(self):
        # Folder structure
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S_")
        new_dir = os.path.join(self.base_dir, timestamp + "filtered")
        os.makedirs(new_dir)
        new_sub_dir = os.path.join(new_dir, "Passed")
        os.makedirs(new_sub_dir)
        # Filters .txt file
        new_filter_path = os.path.join(new_dir, os.path.basename(self.filter_path))
        shutil.copy2(self.filter_path, new_filter_path)
        # Summary sheet
        self._create_summary_sheet(new_dir)
        # Files filtered in subdir
        self.gui.advance(70)
        for passed in self.passed_list:
            new_doc_path = os.path.join(new_sub_dir, os.path.basename(passed))
            shutil.copy2(passed, new_doc_path)
        logging.debug("Docs that passed filters copied to: %s", new_sub_dir)
        self.gui.advance(100)
        self.gui.success()

    def filter_docs(self):
        """Copies the docs that pass the filter into a new folder"""
        self._import_docs()
        self.gui.advance(25)
        self._import_filters()
        self.gui.advance(30)
        for doc in self.docs:
            if doc.passes_filters(self.filters):
                self.passed_list.append(doc.get_path())
        logging.debug("A total of %s docs passed the filter", len(self.passed_list))
        self.gui.advance(35)
        self._copy_passed_docs()


if __name__ == "__main__":
    main()
